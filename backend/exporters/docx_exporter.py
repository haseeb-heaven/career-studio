import io
from docx import Document
from docx.shared import RGBColor
from exporters.base import Exporter
from exporters import register
from models import Profile


BLUE = RGBColor(0x1E, 0x3A, 0x8A)


def _heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = BLUE
    return p


@register("docx")
class DocxExporter(Exporter):
    mime_type = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    extension = "docx"

    def export(self, profile: Profile) -> bytes:
        doc = Document()
        doc.add_heading(profile.full_name, 0)
        contact_parts = list(filter(None, [profile.email, profile.phone, profile.location]))
        doc.add_paragraph(" | ".join(contact_parts))

        if profile.summary:
            _heading(doc, "Summary")
            doc.add_paragraph(profile.summary)

        if profile.skills:
            _heading(doc, "Skills")
            for s in profile.skills:
                doc.add_paragraph(f"{s.name} ({s.category}) — {s.years} yrs", style="List Bullet")

        if profile.experience:
            _heading(doc, "Experience")
            for e in profile.experience:
                _heading(doc, f"{e.role} — {e.company}", level=2)
                doc.add_paragraph(f"{e.start} – {e.end}  {e.location}".strip())
                for b in (e.bullets or []):
                    doc.add_paragraph(b.text, style="List Bullet")

        if profile.projects:
            _heading(doc, "Projects")
            for p in profile.projects:
                _heading(doc, p.name, level=2)
                if p.description:
                    doc.add_paragraph(p.description)
                if p.get_tech():
                    doc.add_paragraph("Tech: " + ", ".join(p.get_tech()))

        if profile.education:
            _heading(doc, "Education")
            for ed in profile.education:
                doc.add_paragraph(f"{ed.degree} {ed.field} — {ed.institution} ({ed.start}–{ed.end})")

        if profile.certifications:
            _heading(doc, "Certifications")
            for c in profile.certifications:
                doc.add_paragraph(f"{c.name} — {c.issuer} ({c.date})", style="List Bullet")

        if profile.availability or profile.compensation:
            _heading(doc, "Availability & Compensation")
            doc.add_paragraph(f"Available: {profile.availability}  |  Compensation: {profile.compensation}")

        buf = io.BytesIO()
        doc.save(buf)
        return buf.getvalue()
