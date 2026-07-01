"""Live resume/CV editor: AI-seeded draft, freeform editing, Overleaf-style
AI suggestions, and export to plain text / Markdown / DOCX / PDF.

Unlike the structured Profile export panel, drafts here are freeform text
(Markdown-ish) so the user can edit wording directly rather than through
per-field forms — closer to Overleaf than to a template engine.
"""
import io
from pathlib import Path

from docx import Document
from docx.shared import RGBColor
from fpdf import FPDF
from fastapi import APIRouter, HTTPException, Depends, Response
from pydantic import BaseModel
from sqlmodel import Session, select

import db
from models import Profile, ResumeDraft, User
from services import activity
from services.ai_service import complete_complex, profile_text_summary
from logger import get_logger
from routers.auth_utils import get_current_user
from routers.profile_router import _check_ownership

logger = get_logger(__name__)
router = APIRouter(prefix="/profiles", tags=["resume-editor"])

_FONTS_DIR = Path(__file__).parent.parent / "assets" / "fonts"
_FONT_REGULAR = _FONTS_DIR / "DejaVuSans.ttf"
_FONT_BOLD = _FONTS_DIR / "DejaVuSans-Bold.ttf"
_BLUE = (30, 58, 138)


def _get_profile(pid: int, session: Session) -> Profile:
    p = session.get(Profile, pid)
    if not p:
        raise HTTPException(status_code=404, detail="Profile not found")
    session.refresh(p)
    return p


def _get_draft(profile_id: int, draft_id: int, session: Session) -> ResumeDraft:
    draft = session.get(ResumeDraft, draft_id)
    if not draft or draft.profile_id != profile_id:
        raise HTTPException(status_code=404, detail="Draft not found")
    return draft


def _draft_out(d: ResumeDraft) -> dict:
    return {
        "id": d.id,
        "title": d.title,
        "content": d.content,
        "created_at": d.created_at.isoformat(),
        "updated_at": d.updated_at.isoformat() if d.updated_at else d.created_at.isoformat(),
    }


# ---------- generate ----------

class GenerateDraftRequest(BaseModel):
    title: str = ""


@router.post("/{profile_id}/resume-drafts/generate")
def generate_draft(profile_id: int, body: GenerateDraftRequest, user: User = Depends(get_current_user)):
    with Session(db.engine) as s:
        p = _get_profile(profile_id, s)
        _check_ownership(s, p, user)
        resume_text = profile_text_summary(p)

    system = (
        "You are an expert resume writer. Rewrite the resume below into a clean, "
        "well-organized Markdown document ready for a human to fine-tune directly. "
        "Use '# Name' for the title, '## Section' for section headings "
        "(Summary, Skills, Experience, Projects, Education, Certifications), "
        "and '- ' bullet points for achievements. Keep all factual content — "
        "do not invent employers, dates, or skills. Return ONLY the Markdown, no commentary."
    )
    try:
        content = complete_complex(system, f"Resume:\n{resume_text}", user_id=user.id)
    except Exception as e:
        logger.error("Resume draft generation failed: %s", e)
        raise HTTPException(status_code=502, detail=f"AI error: {e}")

    with Session(db.engine) as s:
        draft = ResumeDraft(
            profile_id=profile_id,
            title=body.title or "Untitled Draft",
            content=content,
        )
        s.add(draft)
        s.commit()
        s.refresh(draft)
        activity.log_activity("resume_draft", f"generated draft #{draft.id}", profile_id)
        return _draft_out(draft)


# ---------- list / save / delete ----------

@router.get("/{profile_id}/resume-drafts")
def list_drafts(profile_id: int, user: User = Depends(get_current_user)):
    with Session(db.engine) as s:
        p = _get_profile(profile_id, s)
        _check_ownership(s, p, user)
        rows = s.exec(
            select(ResumeDraft)
            .where(ResumeDraft.profile_id == profile_id)
            .order_by(ResumeDraft.updated_at.desc())
        ).all()
        return [_draft_out(r) for r in rows]


class SaveDraftRequest(BaseModel):
    title: str | None = None
    content: str


@router.put("/{profile_id}/resume-drafts/{draft_id}")
def save_draft(profile_id: int, draft_id: int, body: SaveDraftRequest, user: User = Depends(get_current_user)):
    with Session(db.engine) as s:
        p = _get_profile(profile_id, s)
        _check_ownership(s, p, user)
        draft = _get_draft(profile_id, draft_id, s)
        draft.content = body.content
        if body.title is not None:
            draft.title = body.title
        s.add(draft)
        s.commit()
        s.refresh(draft)
        return _draft_out(draft)


@router.delete("/{profile_id}/resume-drafts/{draft_id}")
def delete_draft(profile_id: int, draft_id: int, user: User = Depends(get_current_user)):
    with Session(db.engine) as s:
        p = _get_profile(profile_id, s)
        _check_ownership(s, p, user)
        draft = _get_draft(profile_id, draft_id, s)
        s.delete(draft)
        s.commit()
        return {"ok": True}


# ---------- AI suggestions (Overleaf-style: reviewed, not auto-applied) ----------

@router.post("/{profile_id}/resume-drafts/{draft_id}/suggest")
def suggest_edits(profile_id: int, draft_id: int, user: User = Depends(get_current_user)):
    with Session(db.engine) as s:
        p = _get_profile(profile_id, s)
        _check_ownership(s, p, user)
        draft = _get_draft(profile_id, draft_id, s)
        content = draft.content

    system = (
        "You are an expert resume editor. Review the Markdown resume draft below and "
        "return a JSON object: {\"suggestions\": [string, ...]}. Each string is one "
        "specific, actionable suggestion (stronger verb, quantify an achievement, cut "
        "filler, fix a gap) referencing the exact line or phrase it applies to. "
        "Return 5-8 suggestions. Return ONLY valid JSON, no prose."
    )
    try:
        raw = complete_complex(system, f"Draft:\n{content}", user_id=user.id)
        import json
        raw = raw.strip()
        if raw.startswith("```"):
            raw = raw.split("```", 2)[1]
            if raw.startswith("json"):
                raw = raw[4:]
        result = json.loads(raw)
    except Exception as e:
        logger.error("Resume suggestion failed: %s", e)
        raise HTTPException(status_code=502, detail=f"AI error: {e}")

    activity.log_activity("resume_draft", f"suggestions for draft #{draft_id}", profile_id)
    return {"suggestions": result.get("suggestions", [])}


# ---------- export ----------

def _iter_markdown_lines(content: str):
    for raw_line in content.splitlines():
        line = raw_line.rstrip()
        if line.startswith("# "):
            yield "h1", line[2:].strip()
        elif line.startswith("## "):
            yield "h2", line[3:].strip()
        elif line.startswith(("- ", "* ")):
            yield "bullet", line[2:].strip()
        elif line.strip():
            yield "p", line.strip()
        else:
            yield "blank", ""


def _draft_to_docx(content: str) -> bytes:
    doc = Document()
    for kind, text in _iter_markdown_lines(content):
        if kind == "h1":
            h = doc.add_heading(text, level=0)
            for run in h.runs:
                run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
        elif kind == "h2":
            h = doc.add_heading(text, level=1)
            for run in h.runs:
                run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x8A)
        elif kind == "bullet":
            doc.add_paragraph(text, style="List Bullet")
        elif kind == "p":
            doc.add_paragraph(text)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _draft_to_pdf(content: str) -> bytes:
    if not _FONT_REGULAR.exists() or not _FONT_BOLD.exists():
        raise FileNotFoundError(f"Required font not found in {_FONTS_DIR}")
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=19)
    pdf.add_page()
    pdf.set_margins(19, 19, 19)
    pdf.add_font("DejaVu", style="", fname=str(_FONT_REGULAR))
    pdf.add_font("DejaVu", style="B", fname=str(_FONT_BOLD))

    for kind, text in _iter_markdown_lines(content):
        if kind == "h1":
            pdf.set_font("DejaVu", style="B", size=18)
            pdf.set_text_color(*_BLUE)
            pdf.multi_cell(0, 9, text, new_x="LMARGIN", new_y="NEXT")
        elif kind == "h2":
            pdf.ln(1)
            pdf.set_font("DejaVu", style="B", size=12)
            pdf.set_text_color(*_BLUE)
            pdf.multi_cell(0, 8, text, new_x="LMARGIN", new_y="NEXT")
        elif kind == "bullet":
            pdf.set_font("DejaVu", size=10)
            pdf.set_text_color(20, 20, 20)
            pdf.set_x(pdf.l_margin + 5)
            pdf.multi_cell(0, 6, f"• {text}", new_x="LMARGIN", new_y="NEXT")
        elif kind == "p":
            pdf.set_font("DejaVu", size=10)
            pdf.set_text_color(20, 20, 20)
            pdf.multi_cell(0, 6, text, new_x="LMARGIN", new_y="NEXT")
        else:
            pdf.ln(2)
    return bytes(pdf.output())


_EXPORT_MIME = {
    "txt": "text/plain",
    "md": "text/markdown",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    "pdf": "application/pdf",
}


@router.get("/{profile_id}/resume-drafts/{draft_id}/export/{fmt}")
def export_draft(profile_id: int, draft_id: int, fmt: str, user: User = Depends(get_current_user)):
    if fmt not in _EXPORT_MIME:
        raise HTTPException(status_code=400, detail=f"Unsupported format '{fmt}'. Use one of: {', '.join(_EXPORT_MIME)}")

    with Session(db.engine) as s:
        p = _get_profile(profile_id, s)
        _check_ownership(s, p, user)
        draft = _get_draft(profile_id, draft_id, s)
        content = draft.content
        title = draft.title

    if fmt in ("txt", "md"):
        data = content.encode("utf-8")
    elif fmt == "docx":
        data = _draft_to_docx(content)
    else:
        data = _draft_to_pdf(content)

    import re
    safe_title = re.sub(r"[^\w\-]", "_", title).strip("_") or "resume"
    activity.log_activity("export", f"resume draft #{draft_id} → {fmt}", profile_id)
    return Response(
        content=data,
        media_type=_EXPORT_MIME[fmt],
        headers={"Content-Disposition": f'attachment; filename="{safe_title}.{fmt}"'},
    )
